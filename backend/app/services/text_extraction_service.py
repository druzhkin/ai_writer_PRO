"""
Text extraction service for extracting text from various file formats.
"""

import io
import uuid
from typing import Optional, Dict, Any, Tuple
from pathlib import Path
import PyPDF2
from docx import Document
import chardet

from app.utils.file_utils import is_text_file, is_document_file


class TextExtractionService:
    """Service for extracting text from various file formats."""
    
    def __init__(self):
        self.supported_formats = {
            'text/plain': self._extract_text_plain,
            'application/pdf': self._extract_text_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_text_docx,
            'application/msword': self._extract_text_doc,
            'application/rtf': self._extract_text_rtf
        }
    
    async def extract_text(
        self, 
        content: bytes, 
        mime_type: str, 
        filename: Optional[str] = None
    ) -> Tuple[bool, str, Optional[str], Optional[Dict[str, Any]]]:
        """
        Extract text from file content.
        
        Args:
            content: File content as bytes
            mime_type: MIME type of the file
            filename: Optional filename for context
            
        Returns:
            Tuple of (success, error_message, extracted_text, metadata)
        """
        try:
            # Check if format is supported
            if mime_type not in self.supported_formats:
                return False, f"Unsupported file format: {mime_type}", None, None
            
            # Extract text using appropriate method
            extractor = self.supported_formats[mime_type]
            success, message, text, metadata = await extractor(content, filename)
            
            if not success:
                return False, message, None, None
            
            # Post-process extracted text
            if text:
                text = self._post_process_text(text)
                metadata = metadata or {}
                metadata.update(self._get_text_metadata(text))
            
            return True, "Text extracted successfully", text, metadata
            
        except Exception as e:
            return False, f"Text extraction failed: {str(e)}", None, None
    
    async def _extract_text_plain(
        self, 
        content: bytes, 
        filename: Optional[str] = None
    ) -> Tuple[bool, str, Optional[str], Optional[Dict[str, Any]]]:
        """Extract text from plain text files."""
        try:
            # Detect encoding
            detected = chardet.detect(content)
            encoding = detected.get('encoding', 'utf-8')
            confidence = detected.get('confidence', 0)
            
            # Try to decode with detected encoding
            try:
                text = content.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                # Fallback to utf-8 with error handling
                text = content.decode('utf-8', errors='replace')
                encoding = 'utf-8'
                confidence = 0.5
            
            metadata = {
                'encoding': encoding,
                'encoding_confidence': confidence,
                'extraction_method': 'plain_text'
            }
            
            return True, "Text extracted from plain text file", text, metadata
            
        except Exception as e:
            return False, f"Failed to extract text from plain text file: {str(e)}", None, None
    
    async def _extract_text_pdf(
        self, 
        content: bytes, 
        filename: Optional[str] = None
    ) -> Tuple[bool, str, Optional[str], Optional[Dict[str, Any]]]:
        """Extract text from PDF files."""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Check if PDF is encrypted
            if pdf_reader.is_encrypted:
                return False, "PDF is encrypted and cannot be processed", None, None
            
            text_parts = []
            page_count = len(pdf_reader.pages)
            extracted_pages = 0
            skipped_pages = []
            
            # Handle edge case: empty PDF
            if page_count == 0:
                return False, "PDF contains no pages", None, None
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        # Clean up extracted text
                        cleaned_text = self._clean_pdf_text(page_text)
                        if cleaned_text:
                            text_parts.append(cleaned_text)
                            extracted_pages += 1
                        else:
                            skipped_pages.append(page_num + 1)
                    else:
                        skipped_pages.append(page_num + 1)
                except Exception as e:
                    # Skip problematic pages but continue processing
                    skipped_pages.append(page_num + 1)
                    print(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            # Handle edge case: no extractable text
            if not text_parts:
                return False, "No extractable text content found in PDF", None, None
            
            # Handle edge case: very short extracted text
            text = '\n\n'.join(text_parts)
            if len(text.strip()) < 10:
                return False, "Extracted text is too short to be meaningful", None, None
            
            metadata = {
                'page_count': page_count,
                'extracted_pages': extracted_pages,
                'extraction_method': 'pdf_pypdf2',
                'skipped_pages': len(skipped_pages),
                'skipped_page_numbers': skipped_pages[:10],  # Limit to first 10 for metadata
                'extraction_quality': 'good' if extracted_pages / page_count > 0.8 else 'partial'
            }
            
            return True, f"Text extracted from {extracted_pages}/{page_count} pages", text, metadata
            
        except PyPDF2.errors.PdfReadError as e:
            return False, f"Invalid PDF file: {str(e)}", None, None
        except Exception as e:
            return False, f"Failed to extract text from PDF: {str(e)}", None, None
    
    async def _extract_text_docx(
        self, 
        content: bytes, 
        filename: Optional[str] = None
    ) -> Tuple[bool, str, Optional[str], Optional[Dict[str, Any]]]:
        """Extract text from DOCX files."""
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            
            text_parts = []
            paragraph_count = 0
            table_count = 0
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
                    paragraph_count += 1
            
            # Extract text from tables
            for table in doc.tables:
                table_count += 1
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            if not text_parts:
                return False, "No text content found in DOCX file", None, None
            
            text = '\n\n'.join(text_parts)
            
            metadata = {
                'paragraph_count': paragraph_count,
                'table_count': table_count,
                'extraction_method': 'docx_python_docx'
            }
            
            return True, f"Text extracted from {paragraph_count} paragraphs and {table_count} tables", text, metadata
            
        except Exception as e:
            return False, f"Failed to extract text from DOCX file: {str(e)}", None, None
    
    async def _extract_text_doc(
        self, 
        content: bytes, 
        filename: Optional[str] = None
    ) -> Tuple[bool, str, Optional[str], Optional[Dict[str, Any]]]:
        """Extract text from DOC files (legacy Word format)."""
        try:
            # For .doc files, we would need python-docx2txt or similar
            # For now, return an error suggesting conversion to .docx
            return False, "DOC files are not supported. Please convert to DOCX format.", None, None
            
        except Exception as e:
            return False, f"Failed to extract text from DOC file: {str(e)}", None, None
    
    async def _extract_text_rtf(
        self, 
        content: bytes, 
        filename: Optional[str] = None
    ) -> Tuple[bool, str, Optional[str], Optional[Dict[str, Any]]]:
        """Extract text from RTF files."""
        try:
            # Basic RTF text extraction (removes RTF formatting codes)
            text = content.decode('utf-8', errors='replace')
            
            # Remove RTF formatting codes
            import re
            # Remove RTF control words and groups
            text = re.sub(r'\\[a-z]+\d*\s?', '', text)
            text = re.sub(r'[{}]', '', text)
            text = re.sub(r'\\[^a-z]', '', text)
            
            # Clean up extra whitespace
            text = re.sub(r'\s+', ' ', text)
            text = text.strip()
            
            if not text:
                return False, "No text content found in RTF file", None, None
            
            metadata = {
                'extraction_method': 'rtf_basic',
                'formatting_removed': True
            }
            
            return True, "Text extracted from RTF file", text, metadata
            
        except Exception as e:
            return False, f"Failed to extract text from RTF file: {str(e)}", None, None
    
    def _clean_pdf_text(self, text: str) -> str:
        """
        Clean PDF-extracted text.
        
        Args:
            text: Raw PDF text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        import re
        
        # Remove excessive whitespace and normalize
        text = re.sub(r'\s+', ' ', text)
        
        # Remove common PDF artifacts
        text = re.sub(r'[^\w\s.,!?;:()\-"\']', '', text)
        
        # Remove very short "words" that are likely artifacts
        words = text.split()
        cleaned_words = [word for word in words if len(word) > 1 or word in '.,!?;:()']
        
        return ' '.join(cleaned_words).strip()
    
    def _post_process_text(self, text: str) -> str:
        """
        Post-process extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Processed text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        import re
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Multiple newlines to double
        text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
        text = re.sub(r'\n ', '\n', text)  # Remove leading spaces from lines
        
        # Remove control characters
        text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t')
        
        return text.strip()
    
    def _get_text_metadata(self, text: str) -> Dict[str, Any]:
        """
        Get metadata about extracted text.
        
        Args:
            text: Extracted text
            
        Returns:
            Text metadata
        """
        if not text:
            return {}
        
        # Basic text statistics
        lines = text.split('\n')
        words = text.split()
        characters = len(text)
        characters_no_spaces = len(text.replace(' ', ''))
        
        # Count sentences (rough estimate)
        import re
        sentences = re.split(r'[.!?]+', text)
        sentence_count = len([s for s in sentences if s.strip()])
        
        return {
            'character_count': characters,
            'character_count_no_spaces': characters_no_spaces,
            'word_count': len(words),
            'line_count': len(lines),
            'sentence_count': sentence_count,
            'avg_words_per_sentence': len(words) / sentence_count if sentence_count > 0 else 0,
            'avg_characters_per_word': characters / len(words) if words else 0
        }
    
    async def extract_text_with_fallback(
        self, 
        content: bytes, 
        mime_type: str, 
        filename: Optional[str] = None
    ) -> Tuple[bool, str, Optional[str], Optional[Dict[str, Any]]]:
        """
        Extract text with fallback methods.
        
        Args:
            content: File content as bytes
            mime_type: MIME type of the file
            filename: Optional filename for context
            
        Returns:
            Tuple of (success, error_message, extracted_text, metadata)
        """
        # Try primary extraction method
        success, message, text, metadata = await self.extract_text(content, mime_type, filename)
        
        if success and text:
            return success, message, text, metadata
        
        # Try fallback methods
        if mime_type == 'application/pdf':
            # Try alternative PDF extraction
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(stream=content, filetype="pdf")
                text_parts = []
                for page in doc:
                    text_parts.append(page.get_text())
                doc.close()
                
                if text_parts:
                    text = '\n\n'.join(text_parts)
                    metadata = metadata or {}
                    metadata.update({
                        'extraction_method': 'pdf_pymupdf_fallback',
                        'fallback_used': True
                    })
                    return True, "Text extracted using fallback method", text, metadata
            except ImportError:
                pass
            except Exception:
                pass
        
        return success, message, text, metadata
